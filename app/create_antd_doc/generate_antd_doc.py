import re
import pandas as pd
import numpy as np
import docx
from docx.shared import Inches, Cm
from docx.shared import RGBColor
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from config import TOP_N_RESULTS
import logging

logger = logging.getLogger(__name__)

def float2percent(string):
    import ast
    if isinstance(string, str):
        s = ast.literal_eval(string)
    else:
        s = string
    s = tuple(["{:.0%}".format(i) if i else i for i in s])
    return s

def overall_judge(x):
    judge_lst = x.judge
    manual_identifier = x.manualIdentifier
    match_term_list = x.matchTermList
    
    if judge_lst is None and not manual_identifier and not match_term_list:
        return 'N/A'
    elif judge_lst is None and not manual_identifier:
        return 'All Matched'
    elif judge_lst == False or all(i==False for i in judge_lst) and not manual_identifier:
        return 'N/A'
    elif any(i==True for i in judge_lst) and not manual_identifier:
        return 'All Matched'
    elif all(i==False for i in judge_lst) and manual_identifier:
        return 'All Not Matched'
    elif any(i==True for i in judge_lst) and manual_identifier:
        return 'Partially Matched'

def judge_grp(judge_all_grp):
    '''
    Given a 'judge_all' pd.Series from pd.groupby object
    check if there is any 'TP' in a group of judges, return judge_grp as 'TP'
    elif all 'TN' were found in a group of judges, return judge_grp as 'TN'
    elif all 'FN' were found in a group of judges, return judge_grp as 'FN'
    elif all 'FP' were found in a group of judges, return judge_grp as 'FP'
    else return ['FP','FN'] as there should be some 'FP' and some 'FN' in the group
    e.g. judge_all_grp = pd.Series(['FP','FP','TP','FP','FP']), return 'TP'
    e.g. judge_all_grp = pd.Series(['FP','FP','FN']), return ['FP','FN']
    '''
    if judge_all_grp.eq('TP').any():
        return 'TP'
    elif judge_all_grp.eq('TN').all():
        return 'TN'
    elif judge_all_grp.eq('FN').all():
        return 'FN'
    elif judge_all_grp.eq('FP').all():
        return 'FP'
    elif judge_all_grp.eq('FP').any() and judge_all_grp.eq('FN').any():
        return str(['FP', 'FN'])
    # elif judge_all_grp.eq(None).any() or judge_all_grp.eq('None').any() or judge_all_grp.eq(np.nan).any():
    #     return None, 1/judge_all_grp.size
    else:
        return 'Indeterminate'


class RegenerateWords(object):

    def __init__(self, term_match_data, ts_filename, fa_filename, output_path, with_annot=False):
        self.df = pd.DataFrame(term_match_data)
        self.df2 = None
        self.ts_filename = ts_filename
        self.fa_filename = fa_filename
        self.output_path = output_path
        self.document = docx.Document()
        self.with_annot = with_annot
        self.header = self.document.sections[0].header
        self.footer = self.document.sections[0].footer
        self.last_text = None
        self.last_id_list = None
        self.last_index = None
        self.p = self.document.add_paragraph()
        self.apply_comment_style = False
        self.marginCm = 1.27
        self.table_id = 0
        self.topN = TOP_N_RESULTS
        self.id_below_fifty = dict()
        self.id_with_matched = dict()
        self.map_judgement = {
            "FP": "Irrelavant",
            "TP": "Matched",
            "FN": "False",
            "TN": "N/A",
            'Indeterminate': 'Indeterminate',
            None: "Indeterminate",
            np.nan: "Indeterminate",
            str(['FP', 'FN']): "Indeterminate"
        }

    def append_table(self, data):

        if self.with_annot:
            # Creating a table object
            table = self.document.add_table(rows=1, cols=4)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Matched Id'
            row[1].text = 'Matched Clause Content'
            row[2].text = 'Similairty Score'
            row[3].text = 'Judgement'
        else:
            # Creating a table object
            table = self.document.add_table(rows=1, cols=4)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Matched Id'
            row[1].text = 'Matched Clause Content'
            row[2].text = 'Similairty Score'
            row[3].text = 'Judgement'
            # row[4].text = 'Overall Judgement'

        # self.table_id = len(self.document.tables)

        self.id_below_fifty[self.table_id] = []
        self.id_with_matched[self.table_id] = []
        matched_flag = False
        # Adding data from the list to the table
        for i, item in enumerate(data):
            if self.with_annot:
                Id, term, sim_score, judge = item
                if judge == 'Matched' and sim_score:
                    self.id_with_matched[self.table_id].append(i+1)
                    matched_flag = True
            else:
                Id, term, sim_score, judge = item
            if sim_score and int(sim_score.split('%')[0]) < 50:
                self.id_below_fifty[self.table_id].append(i+1)

            # Adding a row and then adding data in it.
            row = table.add_row().cells
            table_msg = ''
            if term and term.startswith('[') and term.endswith(']'):
                table_msg = '<This is a table>\n'
            # Converting id to string as table can only take string input
            row[0].text = Id.replace('【','\n【') if Id else ''
            row[1].text = table_msg + term if term else ''
            row[2].text = sim_score if sim_score else ''
            if self.with_annot:
                row[3].text = judge if sim_score else 'N/A'
            else:
                row[3].text = str(judge) if sim_score else 'N/A'

        if not self.with_annot:
            colId_widthCm = [(0, 6.44), (1, 9.53), (2, 2.15), (3, 2.15)]
            for col, width in colId_widthCm:
                for cell in table.columns[col].cells:
                    cell.width = Cm(width)

        if not self.with_annot:
            # Adding style to a table
            table.style = 'Light List'
        elif self.with_annot and matched_flag:
            table.style = 'Light List Accent 3'
        else:
            table.style = 'Light List Accent 5'

    def transform2table(self, data):
        import ast
        import json
        if isinstance(data, str) and data.startswith('["'):
            data = ast.literal_eval(data)
        assert isinstance(data, list) and all(isinstance(elem, dict) for elem in data), f'Please assure input table data is a list of dict. data: {json.dumps(data,indent=4)}'
        # Creating a table object
        table = self.document.add_table(rows=len(data)+1, cols=len(list(data[0].keys())), style="Table Grid")

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="AEAEAE"/>'.format(nsdecls('w')))
        for i, k in enumerate(list(data[0].keys())):
            row[i].text = k
            row[i].bold = True
            # row[i]._tc.get_or_add_tcPr().append(shading_elm_1) # fill the i cell in the first row of your table with the RGB color AEAEAE

        # Adding data from the list to the table
        for dic in data:
            # Adding a row and then adding data in it.
            row = table.add_row().cells
            for i, value in enumerate(list(dic.values())):
                # Converting id to string as table can only take string input
                row[i].text = value if value else ''

        table.autofit = True

    def generate(self):
        import ast
        import json

        last_sec = None
        is_first_title = True
        footer_para = self.footer.paragraphs[0]
        footer_para.text = f"TS filename: {self.ts_filename}\nFA filename: {self.fa_filename}"

        # changing the page margins
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(self.marginCm)
            section.bottom_margin = Cm(self.marginCm)
            section.left_margin = Cm(self.marginCm)
            section.right_margin = Cm(self.marginCm)

        self.df = self.df.replace({np.nan: None, 'None': None})
        if self.with_annot:
            self.df['judge'] = self.df['judge'].map(lambda x: tuple(ast.literal_eval(str(x))) if isinstance(x, str) and x.startswith('(') and x != 'None' else x)
            df = self.df.set_index(['indexId', 'fa_identifier'])
            df_gp = df.groupby(level=[0, 1], dropna=False)

            self.df2 = pd.DataFrame(data={
                'tsTerm': df_gp['tsTerm'].transform(lambda x: list(set(x))[0]),
                'tsText': df_gp['tsText'].transform(lambda x: ' '.join(x) if isinstance(x,list) else x),
                'judge_grp': df_gp.judge_all.transform(lambda x: judge_grp(x)),
            }).reset_index(names=['indexId', 'fa_identifier']).drop_duplicates(subset=['indexId', 'fa_identifier'])

            self.df2 = self.df2.replace(np.nan, None)

            df_gp2 = self.df2.groupby('judge_grp', dropna=False)['indexId'].count().transpose()
            df_gp2 = df_gp2.rename({'FN': 'fn', 'FP': 'fp', 'TN': 'tn', 'TP': 'tp'})
            df_gp2 = df_gp2.to_frame().transpose().reset_index(drop=True)
            judges = ['tp', 'fp', 'tn', 'fn', 'Indeterminate']
            not_exists = []
            l = []
            for j in judges:
                if j in df_gp2.columns:
                    l.append(j)
                else:
                    not_exists.append(j)
            df_gp2 = df_gp2[l]
            if not_exists:
                for i in not_exists:
                    df_gp2[i] = 0

            totalMatchedCount = df_gp2.tp.values[0]
            totalFalseCount = df_gp2.fn.values[0]
            totalIrrelevantCount = df_gp2.fp.values[0]
            totalNACount = df_gp2.tn.values[0]
            totalIndeterminateCount = df_gp2.Indeterminate.values[0]
            HitRate = str(round(totalMatchedCount / (totalMatchedCount + totalFalseCount) * 100, 2)) + ' %'
            # Creating a table object
            table = self.document.add_table(rows=1, cols=4)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Matched Count'
            row[1].text = 'False Count'
            row[2].text = 'Irrelevant Count'
            row[3].text = 'Indeterminate Count'
            row[4].text = 'Hit Rate = Matched / (Matched + False)'

            row = table.add_row().cells
            row[0].text = str(totalMatchedCount)
            row[1].text = str(totalFalseCount)
            row[2].text = str(totalIrrelevantCount)
            row[3].text = str(totalIndeterminateCount)
            row[4].text = HitRate

            table.style = 'Light List Accent 1'
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        else:
            self.df.loc[~self.df.matchTermList.isna(),'judge'] = self.df[~self.df.matchTermList.isna()]['result'].map(lambda x: [True]+[False]*(self.topN-1) if x is None else [bool(i) for i in x])
            self.df.loc[~self.df.matchTermList.isna(),'judge_all'] = self.df[~self.df.matchTermList.isna()][['judge','manualIdentifier','matchTermList']].apply(lambda x: overall_judge(x), axis=1)
            self.df.loc[(self.df.manualIdentifier.isna()), 'judge_all'] = 'N/A'

        for _, row in self.df.iterrows():
            textElement = row['textElement']
            tsText = row['tsText']
            listId = row['listId']
            indexId = row['indexId']
            identifierList = row['identifierList']
            matchTermList = row['matchTermList']
            similarityList = row['similarityList']
            pageId = row['pageId']
            if similarityList not in [None, np.nan, '', 'None']:
                similarityList = float2percent(similarityList)
                # similarityList = similarityList[:self.topN]

            if self.last_text == tsText or textElement is None:
                continue

            if not textElement in ['section', 'title']:
                self.p = self.document.add_paragraph()
                sec_name = last_sec + ':' if last_sec and not last_sec.endswith(':') else last_sec
                run = self.p.add_run(f'【 section: {sec_name} 】【 text group ID: {str(indexId)} 】【 text type: {textElement} 】\n')
                run.bold = True
                run.underline = True

            if textElement == 'section':
                h = self.document.add_paragraph(tsText, style='Intense Quote')
                # 0 for left, 1 for center, 2 right, 3 justify ....
                h.alignment = 1
                last_sec = tsText
                # continue
                # self.p = self.document.add_paragraph()
                # run = self.p.add_run(tsText)
                # run.bold = True
            elif 'title' in textElement:
                if is_first_title:
                    header_para = self.header.paragraphs[0]
                    header_para.text = tsText
                    is_first_title = False
                t = self.document.add_heading(tsText, level=1)
                # 0 for left, 1 for center, 2 right, 3 justify ....
                t.alignment = 1
                continue
            elif 'list' in textElement:
                if listId not in [None, np.nan, '', 'None']:
                    c = self.document.add_paragraph()
                    c.add_run(str(listId)+'\t')
                    run = c.add_run(tsText)
                    run.left_indent = Inches(1)
                else:
                    self.document.add_paragraph(tsText, style='List Bullet')
            elif 'caption' in textElement:
                c = self.document.add_paragraph(tsText)
                c.italic = True
            elif 'table' in textElement:
                try:
                    self.transform2table(tsText)
                except AssertionError:
                    self.p = self.document.add_paragraph()
                    run = self.p.add_run(tsText)
            else:
                if identifierList == self.last_id_list and indexId == self.last_index:
                    run = self.p.add_run(tsText)
                    continue
                else:
                    self.p = self.document.add_paragraph()
                    run = self.p.add_run(tsText)

            if self.with_annot:

                judge = row['judge']
                judge_all = row['judge_all'] if row['judge_all'] else None
                fa_identifier = row['fa_identifier'] if row['fa_identifier'] else None
                fa_text = row['fa_text'] if row['fa_text'] else None
                judgement = 'N/A' if identifierList == '(None,)' else self.map_judgement.get(judge_all)
                extra_msg = 'and above text didn\'t apply term matching' if identifierList in ['(None, None, None, None, None, None, None)', '[None]', '(None,)', [None], None] else ''
                self.p = self.document.add_paragraph()
                if fa_identifier:
                    if textElement not in ['section', 'title']:
                        run = self.p.add_run('Annotation of above TS content is: 【 ' + fa_identifier + f' 】{extra_msg}; Overall judgement is 【 ' + judgement + ' 】')
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        self.p = self.document.add_paragraph()
                        run = self.p.add_run('FA content of Annotation:【 ' + fa_text + ' 】' if fa_text else f'FA content cannot be queried. Please check if FA clause ID【 {fa_identifier} 】exists in FA or FA was parsed correctly.')
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    if textElement not in ['section', 'title']:
                        if extra_msg == '':
                            extra_msg = 'but there are term matching results'
                        run = self.p.add_run(f'No annotation is provided for above TS content {extra_msg}; Overall judgement is 【 ' + judgement + ' 】')
                        run.bold = True
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
            else:
                judge = row['judge']
                judge_all = row['judge_all'] if row['judge_all'] else None

            if identifierList not in ['(None, None, None, None, None, None, None)', '[None]', '(None,)', [None], None] and similarityList not in ['(None, None, None, None, None, None, None)', '[None]', '(None,)', [None], None]:
                if isinstance(identifierList, str):
                    identifierList = ast.literal_eval(identifierList)
                    # identifierList = identifierList[:self.topN]
                if isinstance(matchTermList, str):
                    matchTermList = ast.literal_eval(matchTermList)
                    # matchTermList = matchTermList[:self.topN]
                if self.with_annot:
                    if isinstance(judge, str):
                        judge = ast.literal_eval(judge)
                    if judgement == 'Indeterminate':
                        judge = ['Indeterminate'] * self.topN
                    else:
                        judge = judge[:self.topN]
                        judge = [self.map_judgement.get(i) for i in judge]
                # else:
                #     judge = judge[:self.topN]

                results = dict(zip(identifierList, matchTermList))
                result_text = json.dumps(results, indent=2).replace('"', '').replace('{', '').replace('}', '')
                similarityList2 = dict(zip(identifierList, similarityList))
                similarityList_text = json.dumps(similarityList2, indent=2).replace('"', '').replace('{', '').replace('}', '')
                layout_info = {'Term Sheet Page': pageId, f'Term Sheet Section': last_sec}
                layout_info_text = json.dumps(layout_info, indent=2).replace('"', '').replace('{', '').replace('}', '')
                sim_score_text = str(similarityList)

                if self.with_annot:
                    all_results = list(zip(identifierList, matchTermList, similarityList, judge))
                else:
                    all_results = list(zip(identifierList, matchTermList, similarityList, judge))

                if self.apply_comment_style:
                    run.add_comment(layout_info_text, author='This text is extracted from: ', initials='od')
                    run.add_comment(result_text, author='Match FA ID & content: ', initials='od')
                    run.add_comment(sim_score_text, author='Similarity Scores: ', initials='od')
                else:
                    p = self.document.add_paragraph()
                    # run = p.add_run('Above tsText is extracted from: \n')
                    # run.bold = True
                    # run.underline = True
                    # run.font.color.rgb = RGBColor(50, 0, 255)

                    # run = p.add_run(layout_info_text)
                    # run.font.color.rgb = RGBColor(50, 0, 255)
                    if textElement == 'section':
                        run = p.add_run('\nResults matched with above TS Section: ')
                    else:
                        run = p.add_run('\nResults matched with above TS Content: ')
                    run.bold = True
                    run.underline = True

                    self.append_table(all_results)
                    self.table_id += 1
                    self.document.add_paragraph()

                    # run = p.add_run('\nMatch FA ID & content: \n')
                    # run.bold = True

                    # run.underline = True
                    # run.font.color.rgb = RGBColor(0xFF, 0000, 0000)

                    # run = p.add_run('\t'+'\n\t'.join(aligned_results))
                    # run.font.color.rgb = RGBColor(0xFF, 0000, 0000)

                    # run = p.add_run('\nSimilarity Scores: \n')
                    # run.bold = True
                    # run.underline = True
                    # run.font.color.rgb = RGBColor(0xFF, 0000, 0000)

                    # run = p.add_run(sim_score_text)
                    # run.font.color.rgb = RGBColor(0xFF, 0000, 0000)

            self.last_id_list = identifierList
            self.last_index = indexId
            self.last_text = tsText

        if not self.with_annot:
            # for table_id, id_list in self.id_below_fifty.items():
            #     if id_list:
            #         for Id in id_list:
            #             self.document.tables[table_id].rows[Id].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            #             self.document.tables[table_id].rows[Id].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            #             self.document.tables[table_id].rows[Id].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            #             self.document.tables[table_id].rows[Id].cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            # style = self.document.styles['Intense Quote']
            # font = style.font
            # font.size = Pt(13)
            pass
        else:
            for table_id, id_list in self.id_with_matched.items():
                if id_list:
                    for Id in id_list:
                        try:
                            self.document.tables[table_id].rows[Id].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                            self.document.tables[table_id].rows[Id].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                            self.document.tables[table_id].rows[Id].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                            self.document.tables[table_id].rows[Id].cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                            self.document.tables[table_id].rows[Id].cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        except IndexError:
                            logger.error(f'table_id: {table_id}, row_id: {Id}, number of rows: {len(self.document.tables[table_id].rows)}')
            style = self.document.styles['Intense Quote']
            font = style.font
            font.size = Pt(13)

        if self.with_annot:
            self.p = self.document.add_paragraph()
            run = self.p.add_run('Overall performance evaluation based on annotation within a text group:')
            # Creating a table object
            table = self.document.add_table(rows=1, cols=4)
            # Adding heading in the 1st row of the table
            row = table.rows[0].cells
            row[0].text = 'Text Group Index'
            row[1].text = 'Section'
            row[2].text = 'Content'
            row[3].text = 'Annotation'
            row[4].text = 'Judgement within text group'

            for Id, record in self.df2.iterrows():
                # Adding a row and then adding data in it.
                row = table.add_row().cells
                judgement = self.map_judgement.get(record['judge_grp'])
                row[0].text = str(record['indexId']) if record['indexId'] else ''
                row[1].text = record['tsTerm'] if record['tsTerm'] else ''
                row[2].text = record['tsText'] if record['tsText'] else ''
                row[3].text = record['fa_identifier'] if record['fa_identifier'] else 'No annotation'
                row[4].text = judgement if judgement else 'Indeterminate'

                if judgement == 'Matched':
                    row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                elif judgement == 'Indeterminate':
                    row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                elif judgement == 'False':
                    row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
                    row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
                    row[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
                    row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)

            table.style = 'Light List'

        self.document.save(self.output_path)

        return self.output_path


if __name__ == '__main__':
    import os
    import pandas as pd
    import sys
    import inspect
    import warnings
    import argparse
    warnings.simplefilter(action='ignore', category=FutureWarning)
    warnings.simplefilter(action='ignore', category=UserWarning)

    try:
        from config import OUTPUT_ANNOT_DOC_DIR
    except ModuleNotFoundError:
        currentdir = os.path.dirname(os.path.abspath(
            inspect.getfile(inspect.currentframe())))
        parentdir = os.path.dirname(currentdir)
        parentparentdir = os.path.dirname(parentdir)
        sys.path.insert(0, parentparentdir)
        from config import OUTPUT_ANNOT_DOC_DIR

    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--term_match_folder",
        default='/home/data/ldrs_analytics/data/antd_ts_merge_fa/1208_finetune_top5_filter_human_errors/(0)/',
        type=str,
        required=False,
        help="The input term matching result CSV directory.",
    )
    parser.add_argument(
        "--fa_folder",
        default='/home/data/ldrs_analytics/data/docparse_csv/FA_v3.2/',
        type=str,
        required=False,
        help="The input docparse FA CSV directory.",
    )
    parser.add_argument(
        "--output_folder",
        default=OUTPUT_ANNOT_DOC_DIR,
        type=str,
        required=False,
        help="The output annotated term match MS Word document directory.",
    )
    parser.add_argument(
        "--target_list",
        default=[
            "1_GL_SYN_TS_mkd_20221215_docparse_merge&match_FA.csv",
            "13_BL_SYN_TS_mkd_20220713_docparse_merge&match_FA.csv",
            "19_GL_SYN_TS_mkd_20220718_docparse_merge&match_FA.csv",
            "23_BL_SYN_TS_mkd_20220715_docparse_merge&match_FA.csv",
            "24_GF_PRJ_TS_mkd_20220225_docparse_merge&match_FA.csv",
            "25_NBFI_PRJ_TS_mkd_20220613_docparse_merge&match_FA.csv",
            "28_GF_PRJ_TS_mkd_20221111_docparse_merge&match_FA.csv",
            "29_PF_PRJ_TS_mkd_20200624_docparse_merge&match_FA.csv",
            "3_GF_SYN_TS_mkd_20221018_docparse_merge&match_FA.csv",
            "31_GL_PRJ_TS_mkd_20220630_docparse_merge&match_FA.csv",
            "33_BL_PRJ_TS_mkd_20200727_docparse_merge&match_FA.csv",
            "34_AWSP_PRJ_TS_mkd_20211230_docparse_merge&match_FA.csv",
            "41_AF_SYN_TS_mkd_20190307_docparse_merge&match_FA.csv",
            "43_AF_SYN_TS_mkd_20151101_docparse_merge&match_FA.csv",
            "45_AF_PRJ_TS_mkd_20170331_docparse_merge&match_FA.csv",
            "49_VF_PRJ_TS_mkd_20151130_docparse_merge&match_FA.csv",
            "54_VF_PRJ_TS_mkd_20191018_docparse_merge&match_FA.csv",
            "58_VF_SYN_TS_mkd_20111201_docparse_merge&match_FA.csv",
            "59_AWSP_SYN_TS_mkd_20210814_docparse_merge&match_FA.csv",
            "63_AW_SYN_TS_mkd_20221025_docparse_merge&match_FA.csv",
            "66_PF_SYN_TS_mkd_20230106_docparse_merge&match_FA.csv",
            "68_PF_SYN_TS_mkd_20221104_docparse_merge&match_FA.csv",
            "72_NBFI_SYN_TS_mkd_20221215_docparse_merge&match_FA.csv",
            "74_NBFI_SYN_TS_mkd_20220401_docparse_merge&match_FA.csv",
            "8_GF_SYN_TS_mkd_20230215_docparse_merge&match_FA.csv"
            ],
        type=list,
        required=False,
        help="The target term matching CSV files in term_match_folder that used to generate annotated MS document.",
    )
    parser.add_argument(
        "--with_annot",
        default=True,
        type=bool,
        required=False,
        help="True if term matching result CSV contains annotation.",
    )
    args = parser.parse_args()

    if args.target_list:
        if isinstance(args.target_list, str):
            args.target_list = args.target_list.split(',')
        term_match_filenames = [f for f in os.listdir(args.term_match_folder) if f.endswith('.csv') and f in args.target_list]
    else:
        term_match_filenames = [f for f in os.listdir(args.term_match_folder) if f.endswith('.csv')]
    term_match_fpaths = [os.path.join(
        args.term_match_folder, f) for f in term_match_filenames]
    fa_filenames = [re.sub("TS", "FA", f).split('_docparse')[0]+'_docparse.csv' for f in term_match_filenames]
    ts_filenames = [f.split('_docparse')[0]+'_docparse.csv' for f in term_match_filenames]

    fa_name_only = [re.sub('_docparse.csv', '', f) for f in fa_filenames]
    ts_name_only = [re.sub('_docparse.csv', '', f) for f in ts_filenames]

    fa_fpaths = [os.path.join(args.fa_folder, f) for f in fa_filenames]

    all_data = list(zip(term_match_fpaths, fa_fpaths,ts_name_only, fa_name_only))

    for term_match_fpath, fa_fpath, ts_filename, fa_filename in all_data:
        term_match_df = pd.read_csv(term_match_fpath)
        term_match_df = term_match_df.rename(columns={
            'index': 'indexId',
            'text_block_id': 'textBlockId',
            'page_id': 'pageId',
            'phrase_id': 'phraseId',
            'list_id': 'listId',
            'text_element': 'textElement',
            'TS_term': 'tsTerm',
            'TS_text': 'tsText',
            'section': 'tsTerm',
            'text': 'tsText',
            'match_term_list': 'matchTermList',
            'identifier_list': 'identifierList',
            'similarity_list': 'similarityList',
            'match_type_list': 'matchTypeList'
        })

        term_match_data = term_match_df.to_dict('records')
        output_path = os.path.join(args.output_folder, ts_filename + '_matched_results_antd.docx')
        print(f'Generate Annotated Term Match Results for file: {ts_filename}')
        RegenerateWords(term_match_data, ts_filename, fa_filename, output_path, with_annot=args.with_annot).generate()
